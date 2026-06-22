from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT  # FIX: was missing
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

# Title
title = doc.add_heading('THE LAST PLAYSLOT — SEQUENCE OF EVENTS', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.runs[0]
run.font.name = 'Arial'
run.font.size = Pt(16)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 0, 0)

# Subtitle
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('GALA NIGHT | SADMAN HALL | 23 JUNE 2026')
run.font.name = 'Arial'
run.font.size = Pt(12)
run.font.bold = True
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()

# Main Sequence Table
headers = ['Serial', 'Event', 'Performer / Responsible', 'Duration', 'Technical / Remarks']
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_cells = table.rows[0].cells
for i, header in enumerate(headers):
    hdr_cells[i].text = header
    set_cell_shading(hdr_cells[i], 'D9E1F2')
    for paragraph in hdr_cells[i].paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(10)

data = [
    ['—', 'House opens, guests seated', 'Ushers / Reception', '—', 'Background music on loop'],
    ['—', 'Performers standby, sound check', 'Technical Team', '—', 'All mics tested; backup ready'],
    ['1', 'Opening Song', 'Maj', '5 min', 'Photo + lyrics + background visuals on screen'],
    ['—', 'No Filler', '—', '—', 'Direct transition'],
    ['2', 'NEWS of UCSC', 'Maj Nasir', '7 min', 'Script on teleprompter; news desk prop if available'],
    ['3', 'Filler: Meena Cartoon — 3 Wishes', 'AV Team', '5 min', 'Pre-loaded video; sound via PA'],
    ['4', 'Song', 'Maj Anas', '5 min', 'Track + lyrics on screen; mic check before entry'],
    ['5', 'Filler', '', '3 min', 'Standby: solo instrumental / recitation'],
    ['6', 'Demo of DSs', 'DS Panel', '10 min', 'Props / presentation as required; stage lighting'],
    ['—', 'No Filler Required', '—', '—', 'Direct transition'],
    ['7', 'Song', 'Maj Samir &', '6 min', 'Duet track; confirm second name'],
    ['8', 'Filler: DS Reaction', '1 × Offr (absent DS scenario)', '4 min', 'Pre-scripted; comic timing'],
    ['9', 'Song — Keno Ei Nishongota', '', '5 min', 'Track + lyrics; confirm artist'],
    ['10', 'Filler', '', '3 min', 'Standby: short video / musical interlude'],
    ['11', 'Magic Show', 'Mr Fahim', '10 min', 'Table / props setup during prior filler; assistant mic'],
    ['—', 'No Filler Required', '—', '—', 'Direct transition'],
    ['12', 'Stand Up Comedy', 'Mr Mostafa Kabir', '10 min', 'Hand mic; stool / water on stage; spotlight'],
    ['13', 'Demo of DSs', 'DS Panel', '8 min', 'Second demo slot; different content from Serial 6'],
    ['14', 'Tribute to DSs', 'All DSs, OIC, CI, Dy Comdt', '7 min', 'Group on stage; memorial video / photo montage; soft music'],
    ['15', 'Song — E Kon Dorodiya Amar', '', '5 min', 'Track + lyrics; confirm closing artist'],
    ['16', 'General List / Course Events Montage', 'AV Team', '5 min', 'Photo/video compilation of course events'],
    ['17', 'Vote of Thanks', '', '3 min', 'Secretary / OC Gala Night'],
    ['18', 'Memento / Crest Distribution', 'Chief Guest + Host', '5 min', 'Line pre-arranged; 30 sec per presentation'],
    ['19', 'Closing Music & Departure', 'AV Team', '—', 'House lights up; ushers assist exit'],
]

for row_data in data:
    row = table.add_row()
    for i, text in enumerate(row_data):
        cell = row.cells[i]
        cell.text = text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(10)
        if 'Filler' in text and 'No Filler' not in text:
            set_cell_shading(cell, 'FFF2CC')
        elif 'No Filler' in text:
            set_cell_shading(cell, 'F2F2F2')

for row in table.rows:
    row.cells[0].width = Inches(0.7)
    row.cells[1].width = Inches(2.4)
    row.cells[2].width = Inches(2.2)
    row.cells[3].width = Inches(0.9)
    row.cells[4].width = Inches(2.8)

doc.add_paragraph()

# MISSING INFORMATION SECTION
heading = doc.add_heading('MISSING INFORMATION — TO BE FILLED', level=1)
heading.runs[0].font.name = 'Arial'
heading.runs[0].font.size = Pt(13)
heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)

missing_headers = ['Ref', 'Item', 'Action Required']
table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Table Grid'
hdr_cells = table2.rows[0].cells
for i, header in enumerate(missing_headers):
    hdr_cells[i].text = header
    set_cell_shading(hdr_cells[i], 'FFF2CC')
    for paragraph in hdr_cells[i].paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(10)

missing_data = [
    ['A', 'Maj (Opening Song)', 'Confirm name, photo, lyrics, background visuals'],
    ['B', "Maj Samir's duet partner", 'Confirm second name'],
    ['C', 'Keno Ei Nishongota performer', 'Confirm artist'],
    ['D', 'E Kon Dorodiya Amar performer', 'Confirm artist'],
    ['E', 'Filler content (Serial 5, 10)', 'Confirm type: instrumental / recitation / video'],
    ['F', 'Vote of Thanks — speaker', 'Confirm name'],
    ['G', 'Chief Guest', 'Confirm for memento distribution'],
    ['H', 'General List Item 2', 'Complete from draft'],
]

for row_data in missing_data:
    row = table2.add_row()
    for i, text in enumerate(row_data):
        cell = row.cells[i]
        cell.text = text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(10)

for row in table2.rows:
    row.cells[0].width = Inches(0.5)
    row.cells[1].width = Inches(2.5)
    row.cells[2].width = Inches(3.5)

doc.add_paragraph()

# TOP 5 CRITICAL ACTIONS
heading = doc.add_heading('TOP 5 CRITICAL ACTIONS', level=1)
heading.runs[0].font.name = 'Arial'
heading.runs[0].font.size = Pt(13)
heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)

critical_actions = [
    '1. Confirm all blank names immediately — Program print deadline depends on this.',
    '2. Submit all tracks and visuals to AV Team by 22 June 2026 — No last-minute USB handovers.',
    '3. Magic Show props — Pre-stage table during filler at Serial 10 to ensure 2-min changeover.',
    '4. DS Reaction filler — Script to be rehearsed once; comic timing is critical.',
    '5. Tribute to DSs — Pre-brief all DSs on stage formation; photo montage to be cued before they walk on.',
]

for action in critical_actions:
    p = doc.add_paragraph(action, style='List Number')
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(192, 0, 0)

# Footer note
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('— THE LAST PLAYSLOT | SADMAN HALL | 23 JUNE 2026 —\nDISTRIBUTE TO: Program Director, Stage Manager, Backstage Coordinator, Comperes, AV Team —')
run.font.name = 'Arial'
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = RGBColor(128, 128, 128)

output_path = '/home/user/Claude/THE_LAST_PLAYSLOT_Sequence_of_Events.docx'
doc.save(output_path)
print(f"Document saved to: {output_path}")
