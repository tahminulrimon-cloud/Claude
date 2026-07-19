#!/usr/bin/env python3
"""Generate (1) BM's Plan Presentation Brief and (2) Rewritten APRC with doctrinal
notes — ME-BRH OP, DSCSC 2026. Both in DSCSC service-writing style."""
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

NOTE_COLOR = RGBColor(0x1F, 0x4E, 0x79)
GL_COLOR = RGBColor(0x1B, 0x5E, 0x20)
WL_COLOR = RGBColor(0x8B, 0x1A, 0x1A)

def setup(doc):
    st = doc.styles["Normal"]
    st.font.name = "Arial"; st.font.size = Pt(11)
    st.paragraph_format.space_after = Pt(6)
    rpr = st.element.get_or_add_rPr(); rf = rpr.get_or_add_rFonts()
    rf.set(qn("w:eastAsia"), "Arial")
    for s in doc.sections:
        s.left_margin = Inches(1); s.right_margin = Inches(0.8)
        s.top_margin = Inches(0.8); s.bottom_margin = Inches(0.8)

def classification(doc, txt="EXERCISE SECRET"):
    for s in doc.sections:
        for part, align in [(s.header, WD_ALIGN_PARAGRAPH.CENTER), (s.footer, WD_ALIGN_PARAGRAPH.CENTER)]:
            p = part.paragraphs[0]; p.text = ""; p.alignment = align
            r = p.add_run(txt); r.bold = True; r.font.size = Pt(11); r.font.name = "Arial"

def para(doc, text, bold=False, center=False, underline=False, size=11, indent=0,
         space_after=6, italic=False, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold; r.underline = underline; r.italic = italic
    r.font.size = Pt(size); r.font.name = "Arial"
    if color: r.font.color.rgb = color
    return p

def rich(doc, segments, indent=0, space_after=6, justify=True):
    """segments: list of (text, dict(bold, underline, italic, color, size))"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    if indent: p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(space_after)
    for t, fmt in segments:
        r = p.add_run(t)
        r.bold = fmt.get("bold", False); r.underline = fmt.get("underline", False)
        r.italic = fmt.get("italic", False)
        r.font.size = Pt(fmt.get("size", 11)); r.font.name = "Arial"
        if fmt.get("color"): r.font.color.rgb = fmt["color"]
    return p

def heading(doc, text, size=12):
    para(doc, text, bold=True, underline=True, size=size, space_after=8)

def note(doc, text, label="Note (Doctrinal Justification)"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.45)
    p.paragraph_format.space_after = Pt(10)
    r1 = p.add_run(f"{label}.  ")
    r1.bold = True; r1.italic = True; r1.font.size = Pt(10); r1.font.name = "Arial"
    r1.font.color.rgb = NOTE_COLOR
    r2 = p.add_run(text)
    r2.italic = True; r2.font.size = Pt(10); r2.font.name = "Arial"
    r2.font.color.rgb = NOTE_COLOR
    # left border
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "6"); left.set(qn("w:color"), "1F4E79")
    pbdr.append(left); pPr.append(pbdr)
    return p

def glwl(doc, gl, wl, sowhat):
    """GL vs WL doctrinal contrast box."""
    for lab, txt, col in [("GL (BL/own) Doctrine", gl, GL_COLOR),
                          ("WL (RL/SOHB) Tactics", wl, WL_COLOR),
                          ("So What (BM's Deduction)", sowhat, NOTE_COLOR)]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.45)
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"{lab}.  "); r1.bold = True; r1.italic = True
        r1.font.size = Pt(10); r1.font.color.rgb = col; r1.font.name = "Arial"
        r2 = p.add_run(txt); r2.italic = True; r2.font.size = Pt(10)
        r2.font.color.rgb = col; r2.font.name = "Arial"
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "12")
        left.set(qn("w:space"), "6")
        left.set(qn("w:color"), "1B5E20" if col == GL_COLOR else ("8B1A1A" if col == WL_COLOR else "1F4E79"))
        pbdr.append(left); pPr.append(pbdr)
    doc.paragraphs[-1].paragraph_format.space_after = Pt(10)

def table(doc, headers, rows, widths=None, fontsize=9.5, header_fs=9.5):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        r = c.paragraphs[0].add_run(h); r.bold = True
        r.font.size = Pt(header_fs); r.font.name = "Arial"
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "D9E2F3")
        c._tc.get_or_add_tcPr().append(shd)
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            c = t.rows[i + 1].cells[j]; c.text = ""
            r = c.paragraphs[0].add_run(str(cell))
            r.font.size = Pt(fontsize); r.font.name = "Arial"
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

def picture(doc, path, width=9.4, caption=None, landscape_ok=True):
    if not os.path.exists(path):
        para(doc, f"[Graphic placeholder: {path}]", italic=True)
        return
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        para(doc, caption, center=True, italic=True, size=9, space_after=10)

def landscape_section(doc):
    s = doc.add_section()
    new_w, new_h = s.page_height, s.page_width
    s.page_width, s.page_height = new_w, new_h
    from docx.enum.section import WD_ORIENT
    s.orientation = WD_ORIENT.LANDSCAPE
    s.left_margin = s.right_margin = Inches(0.6)
    s.top_margin = s.bottom_margin = Inches(0.6)
    return s

S = os.path.dirname(os.path.abspath(__file__))
M1 = os.path.join(S, "map1_en_disposition.png")
M2 = os.path.join(S, "map2_op_overlay.png")
M3 = os.path.join(S, "sketch_cco_layout.png")

# =====================================================================
# DOC 1 — BM'S PLAN PRESENTATION BRIEF
# =====================================================================
d = Document(); setup(d); classification(d)

para(d, "EXERCISE — FOR ACADEMIC / TRAINING PURPOSES ONLY (DSCSC 2026)", center=True,
     italic=True, size=9, space_after=2)
para(d, "DSCSC — 2026", center=True, bold=True, size=12, space_after=0)
para(d, "ME — BRH OP", center=True, bold=True, underline=True, size=12, space_after=0)
para(d, "PLAN PRESENTATION TO GOC 6 INF DIV", center=True, bold=True, size=13, space_after=0)
para(d, "BRIEF BY BM, 8 INF BDE — BRH OP IN GEN AREA BAUNIA 7140", center=True, bold=True,
     underline=True, size=13, space_after=4)
para(d, "Presented at: HQ 6 Inf Div, SAVAR          On: 120700 Jul 2026          Time Z: FOXTROT",
     center=True, size=10, space_after=2)
para(d, "Ref: Ex Paper ME–BRH OP (Ser 2–4 with Anx); SOHB (WL/RL Tactics); GL Precis — "
        "Aslt Riv Xing, The Attk, Op Appreciation, Orders and Briefings.", center=True,
     size=10, space_after=10)

note(d, "Structure follows Anx E to Ser 4 (A Suggested Plan Pres Format, Ser 1–10). A DSCSC-standard "
        "BM's brief is delivered in first person on behalf of the Comd, is time-disciplined (20–25 min), "
        "runs from situation to plan in a single analytical thread, and — critically — argues every "
        "recommendation from ground, enemy and doctrine rather than asserting it. Each serial below ends "
        "with the doctrinal contrast (GL vs WL) that underpins the argument.", label="Format Note")

heading(d, "OPENING")
para(d, "1.\tGood morning, Sir. I am Maj [student], BM 8 Inf Bde. With your permission, I shall present "
        "Comd 8 Inf Bde's plan for the brH op in gen area BAUNIA 7140. The brief will take 25 mins and will "
        "fol the seq: latest sit and aval res; terr and implications; en's big picture; en's most likely COA; "
        "higher comd's intent; own courses and comparison incl X sites; broad outline plan; br and rafting "
        "plan; Bde and Div X plan in detail incl CCO; and log plan. Maps in use are SHEET 79 D/13 and the "
        "enlargement of gen area BAUNIA, 1 SQ = 1 KM. I shall refer to the live map marking at Anx A, the op "
        "overlay at Anx B and the X area sketch at Anx C. May I proceed, Sir?")

heading(d, "SER 1 — LATEST SIT AND AVAL RESOURCES")
para(d, "2.\tSit En.  RL VII Corps has culminated and its bulk is wdr west, leaving a div (-) to retain "
        "capr terr. In our sector, a rft RL inf bn gp defends the western bank of River DANUBE in gen area "
        "BAUNIA astride the dml rd br. INTSUM 141800F and aerial recce cfm: ext digging along SQ 6741–6939–6739; "
        "MG/auto wpns at GR 681418, 699400 and 715386; ATk wpn at GR 702395; mixed minefds SQ 7236 and 6940, A "
        "pers minefd SQ 6841; tks seen KHULNA 7138, TCVs and tks JOLSHIRI; hy vehs and gun area activity "
        "ISLAMPUR 6438; LP/OPs FERRY GHAT and south VITARA. Arty and mor bombardment continues east of the riv, "
        "and sporadic air raids on 7 Inf Bde posns cont since 08 Jul.")
para(d, "3.\tSit Own.  7 Inf Bde holds the home bank at BAUNIA; 15 Inf Bde is containing a delaying posn at "
        "SATURIA; 25 Inf Div ops on axis SAVAR–CHANDAN–KALIAKAIR–GORAI. 8 Inf Bde, presently concentrated in "
        "area DHAMRAI, is nominated for the aslt riv X.")
para(d, "4.\tResources.  In addition to normal affiliation, the Bde has: Sqn 6 H; one Sp Coy and one ATGM Pl "
        "of 6 Div Sp Bn (UC); two Fd Coy 6 Engr Bn (one INSPUCM); Br Coy 19 Aslt Engr Bn INSPUCM to Div (with "
        "Br Coy 6 Engr Bn — both gp with CCO for br and raft); one SHORAD Bty INSPUCM; 6 Arty Bde less DS regt "
        "and a MLRS Bty in sp; 16 FGA and 3 hel sorties per day; one para inf coy at 30 min NTM (Div res); Bank "
        "Gp of one bn gp fm 7 Inf Bde; CCO formulated on 6 Div Sp Bn (-).")
glwl(d,
 "GL doctrine treats an opposed river X as a deliberate op in four stages — prep, aslt, build-up and "
 "consolidation — demanding cen con of X means, a dedicated Bank Gp to win the fire fight fm the home bank, "
 "and a CCO to regulate call-fwd of veh and tps. Relative superiority of 3:1 at the pt of aslt and engr "
 "effort as the pacing item are the plg gates.",
 "WL (SOHB) def doctrine on a water obs is 'hold fwd — strike early': coys dug in on the bank with mut sp, "
 "obs and fire integrated on likely X sites, a depth coy with a mob det (tk + TCV) tasked for imm C attk "
 "within 1–2 hrs of a foothold, arty pre-registered on X sites and FUPs, and wdr only on orders as terr is a "
 "bargaining chip.",
 "The race is between my build-up and his C attk clock. Every res above is therefore staged against one test: "
 "does it shorten the window between first wave landing and ATk/tk build-up on the far bank?")

heading(d, "SER 2 — TERR AND IMPLICATIONS")
para(d, "5.\tRiver DANUBE, flowing NW–SE at 2–3 m/s, 150–200 m wide and 3–5 m deep, denies fording and "
        "compels a deliberate aslt X. Sandbars between VITARA and FERRY GHAT constrict the channel to about "
        "100 m — halving water exposure but making my likely X sites predictable to RL. The home bank offers "
        "concealed FUPs and boat harbours in BANANI–BARISHAL–AMBAGAN and a rd/tr network adequate for br tn, "
        "though raised rds canalise the CCO flow and demand strict tfc con. The far bank SW of the dml br is "
        "open; movement away fm the riv is dominated by vill and rd junctions. BAUNIA, astride the CL 80 rd "
        "and the dml br, is the decisive terr: its capr collapses the mut sp between VITARA and KHULNA, opens "
        "axis DHAMRAI–SARAIL and unhinges the riv def. SRIPUR gives depth; ISLAMPUR–JOLSHIRI must be cleared "
        "to lift observed arty fire off the X sites and create mnvr space for 9 Armd Bde. X country mob for hy "
        "veh favours the NW of NAIRA and SAMPAN; the SE is soft.")
para(d, "6.\tTerr therefore dictates four imperatives: X where water exposure is shortest; seize BAUNIA "
        "early; expand faster than RL can mount his C attk; and gain depth to PANAM–ISLAMPUR before the Cl 50 "
        "br becomes the CG of the X system.")
glwl(d,
 "GL: site selection is a tac decision taken on the trinity of short exposure, concealed home-bank prep and "
 "far-bank exits onto the decisive obj — engr convenience alone never decides it. Night X with a moonless "
 "4th-qtr sky and rain covers noise and mov.",
 "WL: def economy dictates strength at the obvious sites; SOHB expects RL to accept gaps elsewhere and rely "
 "on obs (LP/OP), pre-planned arty DFs and the mob det to close them.",
 "Hence I attack the site the en can see but cannot hold once suppressed — BANANI — while depicting X at "
 "FERRY GHAT where his LP/OP will faithfully report my deception.")

heading(d, "SER 3 — EN'S BIG PICTURE")
para(d, "7.\tRL's op reqr is time, not terr gain. Having culminated, RL needs the DANUBE line held long "
        "enough to convert occupied terr into diplomatic leverage (UNSC dialogue proposed 30 Jun). The BAUNIA "
        "bn gp is the tac gatekeeper: if it keeps the brH shallow and the Cl 50 br unbuilt, 9 Armd Bde cannot "
        "break out and RL wins its bargaining time. If BAUNIA falls and we reach PANAM–ISLAMPUR, the riv def "
        "is dislocated and RL loses both time and leverage. The bn gp is likely the ME of its bde and will be "
        "rft if the line is breached.")
glwl(d,
 "GL: the appreciation ties en intent to the op level — defeat mechanism must attack the en's purpose (time), "
 "not merely his posns; hence speed is my weapon.",
 "WL: SOHB 'def of a riv line' — hold at all costs fwd, local C attks to restore the bank, fall back on contg "
 "posn (JOLSHIRI) only on orders; destruction of the def force is acceptable to buy time.",
 "My plan must therefore deny time wholesale: simultaneous suppression, night aslt, and armr across before "
 "first lt — not a sequential grind that plays to his clock.")

heading(d, "SER 4 — EN'S MOST LIKELY COA")
para(d, "8.\tCOA A (Most Likely) — Defeat the X at the riv line.  Three coys fwd at VITARA, BAUNIA and "
        "KHULNA, one coy in depth at SRIPUR; contg posn prep at JOLSHIRI; tk hide and TCVs at JOLSHIRI; LP/OPs "
        "N VITARA/S VITARA and FERRY GHAT; gun posn ISLAMPUR; mixed minefds N BAUNIA and S KHULNA — fighting a "
        "decisive def btl by opposing all probable X, then committing tks/TCVs and rft against the first "
        "foothold before rafting matures. Indicators: dml br, deliberate digging, integrated minefds, ATk "
        "sited on the BANANI–BAUNIA exits, tks retained in depth.")
para(d, "9.\tCOA B (Less Likely) — Trade space for time: engage the X with fire, delay fm successive posns "
        "BAUNIA–SRIPUR–JOLSHIRI, disrupt br constr, then wdr before destruction. Less likely because wdr "
        "surrenders the only card RL holds — retained terr.")
para(d, "10.\tDecisive Window.  I assess the decisive period as H to H+6: fm first wave landing until ATk "
        "wpns, F ech and ldg tks are across. RL's local C attk (coy + tk tp fm JOLSHIRI via SRIPUR) can strike "
        "the BAUNIA foothold within 1–2 hrs of ident; his arty at ISLAMPUR can range every X site fm the "
        "outset.")
glwl(d,
 "GL: select en most probable COA on capability + indicators, then design the defeat mechanism against it — "
 "not against the worst case in the abstract.",
 "WL: C attk norms — imm local C attk by depth pl/coy within 30–60 min at coy level, bn-level C attk with "
 "armr within 1–2 hrs, deliberate C attk (bde res) 4–6 hrs; arty concs pre-registered on X sites, FUPs and "
 "far-bank exits; EW to DF own comd nets.",
 "Defeat of COA A pivots on three acts: kill or fix the JOLSHIRI mob det early (FGA + MLRS + ATGM screen "
 "facing SRIPUR–JOLSHIRI), blind the LP/OPs before H, and silence ISLAMPUR guns by CB fm H-15.")

heading(d, "SER 5 — HIGHER COMD'S INTENT")
para(d, "11.\tCAS intends 6 Inf Div to clear axes SAVAR–SATURIA–DAULATPUR and SAVAR–DHAMRAI–BAUNIA–SARAIL–"
        "CHAMPA and evict RL across the IB. GOC 6 Inf Div intends to maint momentum and deny RL time to "
        "consolidate along River DANUBE. 8 Inf Bde is tasked to estb a brH across River DANUBE in gen area "
        "BAUNIA ASP, allowing constr of a Cl 50 br and opening axis DHAMRAI–SARAIL, to facilitate breakout of "
        "9 Armd Bde towards the IB. Key tasks: aslt riv X; destroy en con the selected X area; seize BAUNIA; "
        "expand through SRIPUR; clear KHULNA–JOLSHIRI–ISLAMPUR; protect X sites fm dir and observed indir "
        "fire; secure PANAM–ISLAMPUR line. End state: a secure brH of adequate depth, an op Cl 50 br, "
        "uninterrupted X of fol-up forces and mnvr space for 9 Armd Bde to break out without tac pause.")
note(d, "The intent para in a DSCSC brief is deliberately two-up to one-up: CAS → GOC → own msn, proving "
        "nested purpose. Every phase of the outline plan must trace to one of the GOC's key tasks — the DS "
        "test is 'show me which key task this phase serves.'")

heading(d, "SER 6 — OWN COURSES AND COMPARISON (INCL X SITES)")
para(d, "12.\tComparison of X Sites.", bold=False)
table(d,
 ["Site", "Crossing Value", "Build-up Value", "Threat / Limitation", "Judgment"],
 [["BANANI", "Channel narrowed to ~100 m by sandbars; gentle banks both sides; tr to water, rd imm on far bank",
   "Fastest foothold-to-brH transition; closest to BAUNIA (decisive obj); Cl 50 br site",
   "Predictable — obvious suitability; covered by MG 681418 and mfd SQ 6841",
   "1st pri — main aslt + br site"],
  ["SAMPAN", "Good entry; constricted channel; wh veh fording at places",
   "Disperses aslt waves; second raft stream", "Far bank hard, sandy, moderately steeper; off the decisive axis",
   "2nd pri — alt aslt / raft"],
  ["FERRY GHAT", "Suitable entry/exit, gentle banks", "Rd west fm far bank for dispersal",
   "Deepest water; far fm BAUNIA; LP/OP overwatches; ferry non-op", "3rd pri — deception / contg"],
  ["BAUNIA RD BR", "Excellent alignment, firm banks", "Direct onto decisive axis once secure",
   "Dml pillars obstruct floating br on same site; tactically obvious, covered by fire",
   "Not for initial X — exploit after BAUNIA secure (br poss some dstn away)"],
  ["PURBO PARA / AMBAGAN", "Good entry; narrow channel (PURBO PARA); AMBAGAN gentle banks but deeper water",
   "AMBAGAN feeds 2nd FUP; supports left fwd bn", "PURBO PARA far bank steeper",
   "AMBAGAN — 2nd aslt X site + raft site 2"]],
 widths=[1.0, 1.7, 1.7, 1.9, 1.4], fontsize=8.5)
para(d, "13.\tCourse I.  Aslt at ni with two bns up fm FUPs BANANI and AMBAGAN (19 Inf Bn rt fwd → VITARA; "
        "29 Inf Bn lf fwd → BAUNIA), 39 Inf Bn fol up → SRIPUR in Ph 2; Ph 3 clears KHULNA–JOLSHIRI–ISLAMPUR; "
        "brH line PANAM–ISLAMPUR. Advtg: decisive obj (BAUNIA) capr in Ph 1 by the shortest exposure X; fol-up "
        "bn husbanded as res; max fire sp concentrated on one X area; entry/exits align with objs. Disadv: not "
        "all objs at a go; X at the most obvious site.")
para(d, "14.\tCourse II.  Aslt at ni two bns up fm FUPs BARISHAL–AMBAGAN and SAMPAN → Ph 1 BAUNIA and "
        "KHULNA; Ph 2 SRIPUR–JOLSHIRI; Ph 3 VITARA–ISLAMPUR. Advtg: rational obj per ph; easy dir keeping. "
        "Disadv: splits the aslt astride the riv bend onto two divergent axes; fol-up bn must clear two objs; "
        "SAMPAN far bank steeper delays exit; BAUNIA and KHULNA together in Ph 1 dilutes 3:1 at the decisive pt.")
para(d, "15.\tSelection.  I recommend Course I. It masses two bns and the whole FSP against the decisive obj "
        "in Ph 1, keeps both X sites within one CCO span, and shortens the vulnerable window at the pt where "
        "RL's C attk must come. The obviousness of BANANI is bought off by deception at FERRY GHAT, suppression "
        "of the near MG/mfd belt, and speed in darkness.")
glwl(d,
 "GL attk doctrine: mass at the pt of decision, one strong thrust rather than two weak ones; ME identified "
 "per ph; res constituted and re-constituted each ph.",
 "WL: expects the main X at the best site and plans DFs accordingly — but SOHB also notes RL's rigidity: fire "
 "plans and C attks are triggered by templates and need higher clearance, lagging a fast-moving aslt.",
 "Course I deliberately collides with his template and beats it on tempo: his DFs will fall on a suppressed, "
 "dispersed, night-moving force while my deception feeds his template a second, false X.")

heading(d, "SER 7 — BROAD OUTLINE PLAN (C OF O INCL GP AND FSP)")
para(d, "16.\tMsn.  8 Inf Bde will estb a brH across River DANUBE in gen area BAUNIA 7140 ASP to facilitate "
        "breakout of 9 Armd Bde.")
para(d, "17.\tScheme of Mnvr — 3 Phs (Op Overlay Anx B).")
para(d, "a.\tPh 1 (H Hr 132000 Jul).  Aslt in waves fm BANANI and AMBAGAN; 19 Inf Bn clears VITARA, 29 Inf "
        "Bn clears BAUNIA, secure PL IRON FIST. Purpose: eliminate dir fire threat to X sites. ME: 19 Inf Bn. "
        "End state: VITARA–BAUNIA clear; ISR ferrying ATk/F ech; hy raft constr started.", indent=0.3)
para(d, "b.\tPh 2 (A Hr NB 140030).  39 Inf Bn (+ Sqn 6 H (-) by raft) clears SRIPUR, secures PL TOP GUN. "
        "Purpose: deny dir obsn of X sites; protect br constr fm observed fire. ME: 39 Inf Bn. End state: "
        "SRIPUR clear; Cl 50 br constr started at BANANI.", indent=0.3)
para(d, "c.\tPh 3 (B Hr NB 140500).  19 and 29 Inf Bn (each + tk tp, comp pl, ATGM) clear ISLAMPUR–JOLSHIRI "
        "and KHULNA, reorg on PL HARD PUNCH (brH line PANAM–ISLAMPUR). Purpose: lift sustained indir fire off "
        "the X; create mnvr space. ME: 29 Inf Bn. End state: brH secure; 9 Armd Bde build-up; breakout NB "
        "140900.", indent=0.3)
para(d, "18.\tGp.  19 Inf Bn: one Comp Pl 6 Div Sp Bn, one ATGM Pl less sec. 29 Inf Bn: one Comp Pl, one "
        "ATGM sec. 39 Inf Bn (Ph 2): Sqn 6 H less two tps, one Comp Pl, ATGM Pl less sec. Ph 3: tk tp + comp "
        "pl + ATGM to each aslt bn. Res: inf coy + Div Sp elm + fd pl (Comd: Coy Comd 39 Bn; loc BARISHAL Ph 1, "
        "BAUNIA Ph 2–3); para inf coy remains Div res at 30 min NTM — dmd for JOLSHIRI blocking or FERRY GHAT "
        "exploitation.")
para(d, "19.\tFSP.  Btlfd shaped fm H-30: FGA (16 sorties) on VITARA, BAUNIA, KHULNA, SRIPUR, JOLSHIRI and "
        "gun area ISLAMPUR; MLRS + Med on CB and JOLSHIRI mob det; 3 x Fd Regt + SP in DS of aslt bns with "
        "SOS/DF on far-bank exits; SHORAD over X sites and br; SATA acquiring ISLAMPUR guns; smoke to blind "
        "LP/OPs S VITARA and FERRY GHAT at H-15; Bank Gp (bn gp ex 7 Bde) wins the dir fire fight fm home bank "
        "and thins out only when brH bites.")
para(d, "20.\tDefeat Mechanism.  Destruction, sequenced: blind (LP/OPs, EW silence, smoke) → suppress (FGA, "
        "arty on fwd coys, CB on ISLAMPUR) → strike (two-bn night aslt onto VITARA–BAUNIA) → dislocate (rapid "
        "SRIPUR–JOLSHIRI clearance denies the C attk BUP) → destroy in detail (armr + ATGM across by raft kill "
        "the mob det fwd of PL HARD PUNCH).")
glwl(d,
 "GL: defeat mechanism must name the mechanism (destruction/dislocation) and the sequence of effects, with "
 "the ME per phase carrying the purpose.",
 "WL: the C attk force is the def's CG at bn level (SOHB); its destruction or fixing collapses the 'hold-fwd' "
 "design because fwd coys have no second act.",
 "Killing the JOLSHIRI tk/TCV det before it crosses PL TOP GUN is the single act on which Ph 2–3 hinge; FGA "
 "pri 1 and the ATGM screen orient there fm H+1.")

heading(d, "SER 8 — BR AND RAFTING PLAN INCL TIMING")
para(d, "21.\tEngr Order of Effort (pacing item).  Two Fd Coy 6 Engr Bn: entry/exit dev at BANANI and "
        "AMBAGAN, mfd breaching with aslt bns, then survivability in reorg. Br Coy 6 Engr Bn + Br Coy 19 Aslt "
        "Engr Bn (both with CCO): rafts then br.")
table(d, ["Ser", "Task", "Site", "Start", "Complete (planned)", "Rmk"],
 [["1", "Aslt boats (waves 1–3, 2 bns)", "BANANI + AMBAGAN", "H (132000)", "H+40 min first wave landed",
   "Boats collected BOLP MORAPARA/BOALI fm 131730"],
  ["2", "ISR (inf sp rafts) — ATk wpn, F ech", "Both X sites", "H+60 (2100)", "Cont",
   "Starts imm after 1st wave lands"],
  ["3", "Cl 40/50 hy rafts x 2", "S BANANI + AMBAGAN", "A Hr (140030)", "A+2 (0230) first tk across",
   "Sqn 6 H (-) by raft; comp pl vehs follow"],
  ["4", "Cl 50 floating br (150–200 m)", "BANANI (dstn fm dml br)", "Ph 2 secure (~0430)", "0900–1000 (6–8 hrs constr)",
   "Both br coys; SHORAD + smoke overhead; dml br site rejected (pillars obstruct)"],
  ["5", "Alt br site recce/prep", "AMBAGAN", "Concurrent", "—", "Contg if BANANI interdicted"]],
 widths=[0.4, 2.1, 1.5, 1.1, 1.6, 2.0], fontsize=8.5)
para(d, "22.\tAddl Reqr fwd to Div: one addl br coy eqpt (~60 m reserve bays) to cover current loss and the "
        "AMBAGAN contg; assessed shortfall in aslt boats for simultaneous 2-bn waves met by pooling Div boat "
        "res with CCO by 131700.")
glwl(d,
 "GL: br constr begins only when the site is free of observed dir fire and sustained observed indir fire — "
 "hence br start is tied to Ph 2 (SRIPUR) not to a clock time; rafts bridge the armr gap in between.",
 "WL: arty will re-register on a detected br line within 30–60 min; RL doctrine ranks br sites above tps as "
 "arty tgts.",
 "Therefore: CB wins the br, not the br coys. MLRS/Med CB program peaks at A Hr and again at br start; smoke "
 "and SHORAD are permanent over BANANI fm A Hr.")

heading(d, "SER 9 — BDE AND DIV X PLAN IN DETAIL (CCO)")
para(d, "23.\tCCO (Sketch at Anx C).  CCO on 6 Div Sp Bn (-): CCO HQ with Div tac; X area comd — CO 6 Engr "
        "Bn (tech con) with Bde rep; Bank Master Gp (one pl per X site, ex aslt bns) con boat stations, wave "
        "disciplines and site tfc; Waterway Con Line on the home bank tr line; TCPs 1–4 on CL 80 A2 (DHAMRAI–"
        "BOALI) and CL 60 A1 (PABNA–MORAPARA); call-fwd areas: Bde Assy A DHAMRAI → Bn Assy A KONABARI/"
        "MOHAKHALI/RATNA → A Veh WA PABNA / B Veh WA MORAPARA–BOALI → BOLP MORAPARA and BOALI → FUP. Comm: CCO "
        "net (VHF) + line to br site; nicknames per site; one-way tfc, dispersal 100 m, no lights, MP dets at "
        "TCPs.")
para(d, "24.\tVeh Pri Table (call fwd by CCO).")
table(d, ["Pri", "Serial", "Means", "Timing"],
 [["1", "ATk/ATGM wpn vehs + mor F ech of aslt bns", "ISR/lt raft", "Fm H+60"],
  ["2", "Tks Sqn 6 H (-) + OP/FOO vehs", "Hy raft", "Fm A Hr"],
  ["3", "Comp Pl vehs, amb, sig veh, CP vehs", "Hy raft", "A+1 onwards"],
  ["4", "39 Bn F ech + res coy vehs", "Raft/br", "On br open"],
  ["5", "9 Armd Bde ldg elms (build-up)", "Cl 50 br", "NB 140900 (B+4)"],
  ["6", "A ech, log pkts (Fwd BAA pull)", "Br, by CCO call fwd", "After breakout starts"]],
 widths=[0.5, 3.2, 1.6, 2.0], fontsize=9)
para(d, "25.\tWave / Flight Plan (each aslt bn, 3 waves; boat station drill per GL pam).  Wave 1: 2 x aslt "
        "coy (ldg coys) + engr breach parties + FOOs — silent paddling, H hr. Wave 2: remaining rifle coys + "
        "bn tac + MMG/mor det — H+20. Wave 3: sp coy bal + adm elm — H+40, motorised boats once surprise "
        "lifted. Bank Master Gp recycles boats; reserve boats 15% held at BOLP.")
para(d, "26.\tC2.  Bde Tac with 29 Bn axis fm H+2 (far bank at A Hr); Bde Main DHAMRAI; Div CCO overlay "
        "issued as Anx to op O; success/failure criteria per site (red/green shuttle codes) to CCO net; EW "
        "silence till H, then deception traffic continues on FERRY GHAT feint net.")
glwl(d,
 "GL: the CCO is a comd instrument, not an engr convenience — it must be able to switch the point of main "
 "effort between sites without new orders (pre-designated alt priorities).",
 "WL: RL EW will DF comd nets and cue arty; SOHB expects RL recce patrols/swimmers against boat harbours the "
 "night before.",
 "Hence radio silence to H, line + LO fwd of WCL, boat harbours guarded by Bank Gp standing patrols fm "
 "121900, and the FERRY GHAT feint runs live traffic to give the DF bearings a false centre of mass.")

heading(d, "SER 10 — LOG PLAN")
para(d, "27.\tCen con, decen execution fm Fwd BAA (est DHAMRAI area): 1st line scales carried complete; 2 "
        "DOS amn dumped at gun areas pre-H; boat/raft POL at BOLP; ADS at BOALI with sec at BANANI home bank; "
        "amb relay posts at both X sites, hel CASEVAC LZ at BARISHAL (3 hel sorties); EME lt repair at B Veh "
        "WA; water pt BOALI. Amn pri: arty CB and smoke. Replen across river only on br open except emergency "
        "raft pkts under CCO pri 6.")
para(d, "28.\tRisks.  (a) Daytime mov Assy A → FUP: mitigated by ni mov, AD umbrella, CB. (b) Frontal aslt "
        "through en surface: mitigated by coord FSP, tempo, darkness. (c) Br interdiction: mitigated by CB "
        "peak at br start, smoke/SHORAD, alt site AMBAGAN. Residual risk accepted, Sir, against the GOC's "
        "demand for momentum.")
para(d, "29.\tThat completes my presentation, Sir. Subject to your approval of Course I, warning O is ready "
        "for issue and the Comd's recce completes by 121630. May I have your decision and any modifications, "
        "Sir?")

# Annex pages (landscape) with maps
landscape_section(d); classification(d)
para(d, "ANX A TO BM'S BRIEF — LIVE MAP MARKING: RL DISPOSITION (NATO/APP-6 SYMBOLS)", bold=True,
     underline=True, center=True, size=12)
picture(d, M1, width=9.6,
        caption="Live map marking — every reported/assessed RL loc plotted with APP-6 hostile symbology on the 1 SQ = 1 KM enlargement.")
d.add_page_break()
para(d, "ANX B TO BM'S BRIEF — OP OVERLAY: BRH OP (COURSE I)", bold=True, underline=True,
     center=True, size=12)
picture(d, M2, width=9.6,
        caption="Op overlay — FUPs, X/br/raft sites, aslt axes by ph, PLs IRON FIST / TOP GUN / HARD PUNCH, CCO areas, 9 Armd Bde breakout.")
d.add_page_break()
para(d, "ANX C TO BM'S BRIEF — SKETCH: X AREA / CCO / BANK MASTER GP LAYOUT", bold=True,
     underline=True, center=True, size=12)
picture(d, M3, width=9.6,
        caption="Schematic (not to scale) — GL doctrinal template of the X area applied to BANANI–AMBAGAN.")
d.save(os.path.join(S, "BM_Plan_Pres_Brief_BrH_BAUNIA_DSCSC2026.docx"))
print("doc1 saved")

# =====================================================================
# DOC 2 — REWRITTEN APRC WITH DOCTRINAL NOTES
# =====================================================================
d = Document(); setup(d); classification(d)
para(d, "EXERCISE — FOR ACADEMIC / TRAINING PURPOSES ONLY (DSCSC 2026)", center=True, italic=True,
     size=9, space_after=2)
para(d, "Copy No 1 of 1", size=10, space_after=0)
para(d, "Total Pages: __", size=10, space_after=8)
para(d, "AN APRC ON BRH OP IN GEN AREA BAUNIA 7140", center=True, bold=True, underline=True, size=14,
     space_after=6)
para(d, "For:  GOC, 6 Inf Div\t\t\t\tFmn/Unit: 8 Inf Bde", size=10.5, space_after=0)
para(d, "By:   Brig Gen Reza, Comd 8 Inf Bde", size=10.5, space_after=0)
para(d, "At:   1700 hrs\t\t\t\t\tOn: 07 Jul 2026", size=10.5, space_after=0)
para(d, "Ref:  Ex Paper ME–BRH OP; Map Sheet 79 D/13 and Enlargement Gen Area BAUNIA (1 SQ = 1 KM)",
     size=10.5, space_after=0)
para(d, "Time Z Used Throughout the Aprc: FOXTROT", size=10.5, space_after=10)
note(d, "Rewritten IAW the DSCSC precis 'Operational Appreciation, Orders and Briefings': aim with "
        "limitations; factors each closing in deductions ('so what'); en treated fm his own doctrine (WL/SOHB, "
        "played by RL); courses tested against the en's most probable COA; outline plan traceable to every "
        "deduction. After each para, a Note gives the doctrinal justification — fact → doctrinal implication → "
        "what the en/own force is therefore likely to do.", label="Format Note")

heading(d, "AIM")
para(d, "1.\tTo estb a brH across River DANUBE in gen area BAUNIA 7140 ASP with a view to facilitating "
        "breakout of 9 Armd Bde towards the IB to annihilate the en, with fol limitations:")
para(d, "a.\tA Cl 50 br to be constr in the brH.", indent=0.4, space_after=2)
para(d, "b.\tAxis DHAMRAI–SARAIL to be opened.", indent=0.4)
note(d, "GL: the aim is a single, attainable, mission-verb statement drawn fm the GOC's tasking (Narrative-1) "
        "with limitations that shape — not multiply — the aim. 'ASP' preserves the time imperative the GOC "
        "stressed ('time is the essence'). Both limitations are engr-driven and will later gate site selection "
        "(br constructability) and brH depth (axis opening).")

heading(d, "FACTORS")
para(d, "Grd and Wx", bold=True, underline=True)
para(d, "2.\tGen.  The terr is flat, interspersed with numerous vills, BUAs, cultivated land and water "
        "features. Vills stand 4–6 ft, rds 6–8 ft and trs 3–4 ft above the surroundings. X country mov is poss "
        "except through water-logged and marshy areas; the rainy season has raised water levels. Rivs and "
        "khals flow gen N to S. Four Cl 80 A2, two Cl 50 A2 and two Cl 60 A1 rds serve the sector; all other "
        "rds are Cl 30 and below; brick-soled trs remain motorable but cannot take sustained tfc. Rivers "
        "DHALESHWARI and DANUBE are the maj obs. TONGI, JOYDEBPUR, KALIAKAIR, HEMAYETPUR, GORAI and SAVAR are "
        "the mentionable BUAs. BAUNIA guards the main rd and the (now dml) br over River DANUBE and is "
        "therefore the most imp obj.")
note(d, "Fact: raised rds/vills on a flooded flat plain. Doctrinal implication (GL): mobility corridors are "
        "canalised onto embanked rds — mov tables and tfc con (CCO) become decisive staff work; def potential "
        "of vills means every obj is a QA (quick attack)/deliberate attk problem, not open-country mnvr. WL "
        "implication: SOHB def doctrine anchors coy posns on vills with interlocking MG arcs along linear "
        "obs — expect the en exactly where the APRC later assumes him.")
para(d, "3.\tMet.")
para(d, "a.\tFacts.  First lt 0530; last lt 1830; moon 4th qtr 6th day; wx rainy.", indent=0.4, space_after=2)
para(d, "b.\tFindings.  ~13 hrs daylight, ~11 hrs ni; dark hrs throughout the ni (late, thin moonrise); "
        "rain likely daily.", indent=0.4, space_after=2)
para(d, "c.\tImplications.  (1) Dark hrs mask mov of large bodies incl boat carriage. (2) Rain restricts vis "
        "and makes C2 harder — simple plans, strict nav aids. (3) Rain covers mov noise incl OBM engines. "
        "(4) Current 2–3 m/s with rain: drift allowance for boats/rafts; swimmers/recce parties degraded.",
     indent=0.4)
note(d, "GL: an opposed X seeks surprise as the first principle; a moonless, rainy ni is the classic window — "
        "hence H Hr after full darkness (2000). WL: RL ni surv relies on LP/OPs, para-illum and pre-registered "
        "DFs rather than mass NVD (SOHB); blinding OPs and forcing him to fire unobserved illum wastes his "
        "arty and preserves surprise.")
para(d, "4.\tGTIs/Objs.  BANANI, AMBAGAN, SAMPAN, PURBO PARA, FERRY GHAT, KHULNA, BAUNIA, VITARA, SRIPUR, "
        "ISLAMPUR, JOLSHIRI.")
para(d, "5.\tGTIs Dominating the Waterway.  FERRY GHAT, PURBO PARA, SAMPAN, AMBAGAN, BARISHAL, BANANI, "
        "VITARA, BAUNIA, KHULNA.")
note(d, "GL: GTIs are listed only to be fought over — each either protects a X site, dominates an exit, or "
        "hosts en fire. The far-bank GTIs (VITARA–BAUNIA–KHULNA arc plus SRIPUR–JOLSHIRI–ISLAMPUR depth) "
        "define the minimum brH: the brH line must run beyond en dir fire (Ph 1), dir obsn (Ph 2) and "
        "sustained observed indir fire (Ph 3) of the X sites — this three-ring logic is the doctrinal skeleton "
        "of the phasing.")
para(d, "6.\tApps to River DANUBE.  App 1: SALTA–KONABARI–PABNA–BHOLA–BANANI. App 2: SALTA–KONABARI–PABNA–"
        "BOALI–AMBAGAN. App 3: SALTA–KONABARI–NAIRA–NATORE.")
para(d, "7.\tExits fm River DANUBE.  Exit 1: VITARA–PANAM–NIDAL. Exit 2: BAUNIA–SRIPUR–ISLAMPUR–NIDAL. "
        "Exit 3: BIRULIA–TARABO–ISLAMPUR–NIDAL.")
note(d, "GL: apps are valued by conceal­ment, cl of rd for br tn and separation (one app per aslt bn + one "
        "for CCO); exits by whether they lead onto the decisive obj without lateral mov under fire. App 1/2 "
        "pair with Exit 2 (BAUNIA–SRIPUR) — the axis the aim demands (DHAMRAI–SARAIL). Exit 3 serves the "
        "FERRY GHAT deception story.")
para(d, "8.\tObs.  River DANUBE: 150–200 m wide (narrowing to ~100 m at sandbar constrictions VITARA–FERRY "
        "GHAT), 3–5 m deep, 2–3 m/s — no fording; wet gaps 40–50 ft (canals) with soft beds inland; en mixed "
        "mfds SQ 7236/6940, A pers mfd SQ 6841, mixed mfd KHULNA 7138; improvised obs on poss X sites; dml brs "
        "at BAUNIA and MULADI.")
note(d, "Fact→implication: a 150–200 m, 3–5 m deep gap is an aslt boat + raft + floating br problem (GL engr "
        "tables); the sandbar constriction halves exposure time but is exactly where WL doctrine sites its "
        "obs-fire trap (mfd 6841 + MG 681418 cover the BANANI narrows). Doctrinally the en has told us he "
        "expects us at BANANI — surprise must therefore come fm timing, suppression and deception, not fm "
        "site novelty alone.")
para(d, "9.\tDeductions (Grd and Wx).")
para(d, "a.\tPri of Objs.  BAUNIA astride the main rd and dml br: its capr makes VITARA and KHULNA "
        "untenable and opens the axis; then KHULNA, VITARA, SRIPUR, JOLSHIRI/ISLAMPUR.", indent=0.4, space_after=2)
para(d, "b.\tPri of X Sites.  BANANI (least width, covered prep, shortest tp X time, tr to water, rd on far "
        "bank) → SAMPAN/AMBAGAN (raft capacity, dispersion) → FERRY GHAT (deception/contg).", indent=0.4, space_after=2)
para(d, "c.\tTime of Attk.  At ni (dark, rainy, moonless window).", indent=0.4, space_after=2)
para(d, "d.\tPh of Attk.  Objs dispersed over ~8 km with intervening features — attk in phs.", indent=0.4, space_after=2)
para(d, "e.\tDir of Attk.  Fm the front (east), the only dir the riv line allows; dir keeping simple in "
        "unreccied grd.", indent=0.4, space_after=2)
para(d, "f.\tDeception.  Multiple feasible sites permit depicting X at FERRY GHAT (live LP/OP audience) and "
        "dummy prep at SAMPAN.", indent=0.4)
note(d, "Each deduction is the 'so what' of paras 2–8 in GL appreciation drill. Pri of objs follows the "
        "GL rule that decisive terr = terr whose loss collapses the en's design (BAUNIA breaks WL mut sp "
        "between VITARA and KHULNA and severs the SAVAR–CHAMPA lateral). Phasing flows fm the three-ring brH "
        "logic at para 5's note. Deception exploits a WL known behaviour: LP/OPs report templates, and RL fire "
        "plans re-orient slowly once committed (SOHB).")

para(d, "En", bold=True, underline=True)
para(d, "10.\tAim and Intention.  Imm: defeat/delay any BL X of River DANUBE and retain capr terr. "
        "Ultimate: hold terr as a diplomatic bargaining chip — RL needs time, not ground.")
note(d, "Fm Spl Idea/Narrative-1: VII Corps culminated, wdr under intl pressure, div (-) left to hold. WL "
        "doctrine for a force trading op-level time: deliberate def on the best obs with 'no wdr without "
        "orders'. Likely en conduct: fight fwd hard, C attk early, accept attrition — COA A behaviour.")
para(d, "11.\tComposn.  Rft inf bn gp: inf (4 coys), armr (2 x tk tp), HAT pl, DS arty (fd bty + med bty "
        "assessed), engr (fd pl), EW det.")
para(d, "12.\tDisposn.  Coys at VITARA, BAUNIA (+, astride dml br), KHULNA fwd; depth coy SRIPUR; contg "
        "posn prep JOLSHIRI; tk hide + TCVs JOLSHIRI; guns ISLAMPUR 6438; LP/OPs FERRY GHAT and S VITARA "
        "6643; MG GR 681418/699400/715386; ATk GR 702395; mfds paras 8; digging SQ 6741/6939/6739 "
        "(live map marking at Anx A).")
note(d, "WL bn def norms (SOHB): bn gp frontage on a water obs 6–8 km with three coys fwd on mut sp vills, "
        "one in depth 2–3 km back covering the main axis; tks held as mob det at the depth loc; HAT ambush "
        "sited on the armd app (GR 702395 covers the BANANI–BAUNIA exit onto the rd). Every reported fact "
        "fits the template — high-confidence disposn, which is itself an indicator of COA A (deliberate "
        "'hold-fwd' def, not a delaying screen, which would show fewer mines and no depth digging).")
para(d, "13.\tMorale.  Low — culminated offn, hy cas, home criticism, wdr of parent fmn; but dug-in tps "
        "with a survival stake fight stubbornly fm prep posns.")
note(d, "GL int drill: morale is rated against behaviour, not sympathy — low strategic morale reduces "
        "C attk élan and ni patrol aggressiveness (opportunities for my recce and deception), but does not "
        "soften prepared def (fire fm trenches needs little morale). Plan takes the first, banks nothing on "
        "the second.")
para(d, "14.\tRes and Rft.  Aval — bde-level res within the RL div (-); rft assessed one bn within 12–24 "
        "hrs of the riv line being breached.")
note(d, "Spl Idea: 'this bn is likely the ME of the RL bde and likely to receive rft.' WL: rft flows to the "
        "pt of penetration once higher HQ confirms it — another argument for tempo: the brH must be past PL "
        "TOP GUN before RL divisional wheels turn.")
para(d, "15.\tAir.  Sporadic RL raids on fwd bde posns since 08 Jul; BLAF holds a favourable air sit (16 "
        "FGA + 3 hel sorties/day to the Bde).")
note(d, "GL: 'favourable' ≠ supremacy — daytime concentrations (veh WAs, br line) still need AD (SHORAD "
        "INSPUCM) and dispersal; ni X further discounts the RL air threat at the decisive hrs.")
para(d, "16.\tReasonable Assumptions.  a. Dug-in posns seen at VITARA, BAUNIA, KHULNA, SRIPUR → coy def "
        "posns there. b. Tks at KHULNA and ATk at BAUNIA → RL ATk def biased south. c. Mfds N BAUNIA, VITARA, "
        "S KHULNA → protective mfds all around def locs. d. RL fd arty bty likely ISLAMPUR (hy vehs 6438, "
        "gun-area digging SQ 6937).")
note(d, "Assumption drill (GL): each assumption is anchored to an observed fact and would be invalidated by "
        "a nameable indicator — tasked to recce/SATA accordingly. The southern ATk bias (assumption b) is "
        "exploitable: the BANANI–VITARA (northern) X faces the thinner ATk arc in Ph 1, before tks matter.")
para(d, "17.\tPreferred Tac (WL/SOHB, played by RL).  a. Contg posn prep at JOLSHIRI. b. Stiff resistance, "
        "fight till last; recapture lost GTIs by imm local C attk, or C attk basing on sit. c. Fwd def with "
        "MG/ATk, mfd, concertina and EW — designed to defeat the X at the water's edge.")
note(d, "SOHB WL water-line def in one line: detect (LP/OP, EW) → engage on the water (MG, arty DFs) → "
        "contain the foothold (mfds, ATk) → destroy it (coy/bn C attk with tks fm depth within 1–2 hrs) → if "
        "penetration exceeds bn capacity, fall back on the contg posn (JOLSHIRI) and hold for bde C attk. My "
        "plan must break the chain at 'detect' (deception, ni, smoke, EW silence) and at 'destroy' (kill the "
        "mob det early), and must reach JOLSHIRI before it is manned as a coherent contg posn — the doctrinal "
        "argument for Ph 3's scope and the earliest poss B Hr.")
para(d, "18.\tDeductions (En).  a. En arty (ISLAMPUR) can interfere with FUPs and all X sites → CB by MLRS "
        "and med arty fm H-15, SATA to acquire, peak fires at A Hr and br start. b. En tks with C attk role → "
        "X ATk wpns and F ech vehs imm after the first wave (ISR fm H+60), ATGM screen oriented SRIPUR–"
        "JOLSHIRI. c. En EW active → radio silence till H, line fwd of WCL, feint traffic on FERRY GHAT net. "
        "d. LP/OPs at FERRY GHAT and S VITARA → blind by smoke/fire at H-15; feed FERRY GHAT OP the deception. "
        "e. All previous deductions cfm.")
note(d, "Every en deduction converts a WL behaviour into an own-force task with a timing — the GL test of a "
        "usable appreciation. The C attk window (1–2 hrs) sets the single most important number in the plan: "
        "ATk means must be on the far bank before H+2.")

para(d, "Relative Str", bold=True, underline=True)
para(d, "19.\tTabulated below (ORBAT + Narrative-1 gp):")
table(d, ["Ser", "Item", "Own", "En", "Ratio", "Rmk"],
 [["1", "Inf", "12 x coy (3 bns) + para coy (Div res)", "4 x coy", "3:1", "Aslt norm met at bde level"],
  ["2", "Armr", "4 x tk tp (Sqn 6 H)", "2 x tk tp", "2:1", "Own tks X by raft fm A Hr"],
  ["3", "Arty", "6 x fd bty, 3 x SP bty, 3 x med bty, 1 x MLRS bty", "1 x fd bty, 1 x med bty (assessed)",
    "6:1", "+ SHORAD bty, SATA bty"],
  ["4", "Engr", "2 x fd coy + 2 x br coy (Div INSPUCM), raft & aslt pl", "1 x fd pl", "6:1",
    "Pacing arm of the op"],
  ["5", "ATGM/ATk", "1 x ATGM pl (+ bank gp ATk)", "1 x HAT pl", "1:1", "Parity — hence early X pri 1"],
  ["6", "Air", "16 x FGA + 3 x hel/day", "Sporadic raids", "Fav", "AD umbrella still reqr"]],
 widths=[0.4, 1.0, 2.6, 1.9, 0.7, 1.6], fontsize=8.5)
para(d, "20.\tDeductions (Relative Str).  a. ATk parity → seek bank gp ATk fire fm home bank and X own ATGM "
        "at the fastest opportune moment (ISR/rafts). b. Raft and aslt pl + boat scales → aslt with 2 bns up "
        "is sustainable. c. Two fd coys → dev entry/exit routes at both X sites concurrently. d. 6:1 arty → "
        "afford simultaneous DF sp, CB and deception fire missions. e. All previous deductions cfm.")
note(d, "GL force-ratio doctrine: 3:1 inf at the pt of aslt is the min for a deliberate attk on a prep posn; "
        "the X itself halves combat power until build-up, so the true insurance is fire superiority (6:1 arty) "
        "and tempo. WL will read the same arithmetic: expecting a 2-bn-up X, his arty will fire on both FUP "
        "areas — dispersal, digging in the FUP and short occupancy (<30 min) are therefore drills, not "
        "options.")

para(d, "Assessment of Tasks", bold=True, underline=True)
para(d, "21.\tTabulated below:")
table(d, ["Ser", "Task", "Tps Reqr", "Tps Alot", "Rmk"],
 [["1", "Capr BAUNIA (Ph 1)", "4 x coy, 1 x tk tp", "29 Inf Bn (+): comp pl, ATGM sec; tks join later",
   "ME axis"],
  ["2", "Capr VITARA (Ph 1)", "4 x coy, 1 x tk tp", "19 Inf Bn (+): comp pl, ATGM pl (-); tks join later", ""],
  ["3", "Capr SRIPUR (Ph 2)", "4 x coy, 1 x tk tp", "39 Inf Bn (+): Sqn 6 H (-), comp pl, ATGM pl (-)",
   "Tks by raft"],
  ["4", "Capr KHULNA, JOLSHIRI, ISLAMPUR (Ph 3)", "8 x coy, 1 x tk tp", "19 + 29 Inf Bn, tk tp + comp pl each",
   "After reorg on PL TOP GUN"],
  ["5", "Bank Gp", "4 x coy", "1 x bn gp ex 7 Inf Bde", "Dmd cfm by Div"],
  ["6", "Bank Master", "1 x pl per site", "Ex aslt bns", "Under CCO"],
  ["7", "CCO", "1 x coy +", "6 Div Sp Bn (-) + 1 x inf coy (dmd to Div)", "Incl TCPs, call-fwd areas"],
  ["8", "Indir fire", "3 x fd regt, 1 x med regt, 1 x MLRS bty", "As reqr — gp fm 6 Arty Bde", ""],
  ["9", "AD", "1 x SHORAD bty", "1 x SHORAD bty (INSPUCM)", "Over X sites/br"],
  ["10", "Engr", "2 x fd coy, 2 x br coy", "2 x fd coy; 2 x br coy with CCO", ""],
  ["11", "Res", "1 x coy + armr", "1 x inf coy, sp coy (-), 1 x fd pl; para coy Div res 30 min NTM", ""]],
 widths=[0.4, 2.0, 1.9, 2.6, 1.3], fontsize=8)
para(d, "22.\tDeduction.  Tps reqr (23 x coy eqvt) exceed own 12 x coy — hence: bank gp found by 7 Inf Bde "
        "(one bn gp); CCO manpower demanded fm Div (6 Div Sp Bn (-) + one inf coy); objs taken sequentially in "
        "phs with the fol-up bn recycled; res reconstituted each ph fm the uncommitted bn ('be prep' msns).")
note(d, "GL troops-to-task drill legitimises phasing: when tasks exceed tps by ~2:1, doctrine offers only "
        "three levers — economy (bank gp fm a holding fmn), sequencing (phs), and higher-fmn augmentation "
        "(CCO, br coys, para coy). All three are used; none of the aslt power is mortgaged to rear-area "
        "tasks — a WL gunner counting coys on the water will still meet 3:1 at each obj in turn.")

para(d, "Time and Space", bold=True, underline=True)
para(d, "23.\tFwd time cal (earliest H Hr):")
table(d, ["Ser", "Event", "Time Reqr", "Running Time"],
 [["1", "Time now", "—", "071700 Jul"],
  ["2", "Initial plg, map recce", "12 hrs 30 min", "071700–080530"],
  ["3", "Comd's recce + plan pres to GOC", "~4 days (incl this pres 120700)", "080530–121630"],
  ["4", "Bde O Gp", "2 hrs", "121630–121830"],
  ["5", "Marrying up, prep, boat trg", "12 hrs", "121830–130630"],
  ["6", "Bn/coy comd recce, plg, orders", "6 hrs", "130630–131230"],
  ["7", "Mov to Bde Assy A, final prep", "4 hrs", "131230–131630"],
  ["8", "Mov to Bn Assy A", "1 hr", "131630–131730"],
  ["9", "Collect boats (BOLP), mov to waterway", "2 hrs", "131730–131930"],
  ["10", "Occupy FUP", "30 min", "131930–132000"],
  ["11", "H HR (EARLIEST)", "—", "132000 JUL"],
  ["12", "Clear Ph 1 objs (ISR fm 2100)", "4 hrs", "132000–140001"],
  ["13", "Consolidation / prep", "30 min", "140001–140030"],
  ["14", "A HR (hy raft starts)", "—", "NB 140030"],
  ["15", "Clear Ph 2 (br constr starts)", "4 hrs", "140030–140430"],
  ["16", "Consolidation / prep", "30 min", "140430–140500"],
  ["17", "B HR", "—", "NB 140500"],
  ["18", "Clear Ph 3, expand to brH line", "4 hrs", "140500–140900"],
  ["19", "9 Armd Bde breakout", "—", "NB 140900"]],
 widths=[0.5, 3.4, 1.6, 1.9], fontsize=8.5)
para(d, "24.\tDeductions (T&S).  a. Earliest H Hr 132000 Jul — full darkness, ~11 hrs of ni to fight Ph 1–2 "
        "and start the br before first lt (0530). b. Ph 3 runs into daylight → FGA umbrella and smoke "
        "programmed for 0500–0900. c. Br (6–8 hrs) completes ~1030–1230 under Ph-3 protection — brH line must "
        "be secure NLT 0900, hence B Hr cannot slip past 0500.")
note(d, "GL backwards-planning: the br is the clock. Every timing above is derived by working back fm 'br "
        "open before RL rft (12–24 hrs) arrives' and fwd fm boat availability. WL: expect RL C attks at "
        "first lt (doctrinal preferred hr) — Ph 3's 0500 start deliberately pre-empts a dawn C attk by "
        "hitting JOLSHIRI first.")

heading(d, "EN COURSES")
para(d, "25.\tCourse A.  Def with three coys up at VITARA, BAUNIA and KHULNA along River DANUBE and one coy "
        "in depth at SRIPUR; contg posn JOLSHIRI; tk hide JOLSHIRI; LP/OPs N/S VITARA and FERRY GHAT; guns "
        "ISLAMPUR; mixed mfds N BAUNIA and S KHULNA — fighting a decisive def btl opposing all probable X.")
para(d, "a.\tLikelihood.  Most likely.", indent=0.4, space_after=2)
para(d, "b.\tEff.  (1) Gds the most likely X sites. (2) Def depth compels own tps to clear in phs. (3) Main "
        "rd SAVAR–CHAMPA physically gd by two coys. (4) Res/rft inducted once the riv line is breached or a "
        "foothold estb — C attk against the immature brH is the decisive act.", indent=0.4)
note(d, "Indicator-matched to WL doctrine (para 17 note): digging, mfds, ATk siting, tk retention in depth "
        "and dml br all evidence 'hold fwd–strike early', and RL's strategic need (time/leverage) rewards it. "
        "Therefore treat COA A as the plg baseline, and weight the defeat mechanism against the C attk force, "
        "not the fwd coys alone.")
para(d, "26.\tCourse B.  Hold until the attk is launched; delay, harass, bring rft, conc arty on X sites, "
        "disorg the attk, then wdr fm successive posns preserving combat power.")
para(d, "a.\tLikelihood.  Less likely.", indent=0.4, space_after=2)
para(d, "b.\tEff.  Fewer own cas; area capr quicker; reorg easier; but en escapes destruction and re-forms "
        "on the next line.", indent=0.4)
note(d, "WL resorts to mob def/delay only when preserving force outweighs holding grd — here terr IS the "
        "bargaining chip, so wdr defeats RL's own purpose. Discount but not dismiss: if COA B materialises "
        "(indicator: thinning fires, veh mov west fm SRIPUR by ni), the para coy and armr exploit to PANAM–"
        "ISLAMPUR ahead of the phs — the plan's branches cover both.")
para(d, "27.\tEn's Most Probable COA.  Course A.")

heading(d, "OWN COURSES")
para(d, "28.\tCourse I.  Aslt across River DANUBE in gen area BAUNIA at ni, two bns up (FUPs BANANI and "
        "AMBAGAN; Bde Assy A DHAMRAI; Bn Assy A KONABARI, MOHAKHALI, RATNA; VMA N DHAMRAI; A Veh WA PABNA; B "
        "Veh WA and BOLP MORAPARA and BOALI), estb brH along PANAM–ISLAMPUR in 3 phs: Ph 1 clr VITARA and "
        "BAUNIA to PL IRON FIST; Ph 2 clr SRIPUR to PL TOP GUN; Ph 3 clr KHULNA, JOLSHIRI and ISLAMPUR, reorg "
        "on PL HARD PUNCH.")
para(d, "29.\tAdvtg.  a. Most imp obj (BAUNIA) capr in Ph 1. b. Fol-up bn employable as res. c. Max fire sp "
        "concentrable on one X area. d. Good entry/exits; far-bank exits lead onto the objs. e. Shortest "
        "water exposure (BANANI narrows).")
para(d, "30.\tDisadv.  a. All objs cannot be capr at a go. b. X at the most obvious site — risk of higher "
        "cas without effective suppression/deception.")
para(d, "31.\tCourse II.  As Course I but FUPs BARISHAL–AMBAGAN and SAMPAN, raft sites AMBAGAN and SAMPAN, "
        "br site AMBAGAN: Ph 1 clr BAUNIA and KHULNA (PL IRON DOME); Ph 2 SRIPUR and JOLSHIRI (PL TOM CAT); "
        "Ph 3 VITARA and ISLAMPUR (PL HARD).")
para(d, "32.\tAdvtg.  a. Rational objs each ph. b. Easy dir keeping. c. Good entry/exits at the riv.")
para(d, "33.\tDisadv.  a. Fol-up bn must clr two objs. b. Still X at an obvious site (SAMPAN watched fm "
        "KHULNA). c. Aslt splits astride the riv bend onto divergent axes — mut sp between fwd bns lost. "
        "d. SAMPAN far bank hard/steeper — exit delay under fire. e. VITARA left alive on the north flank of "
        "the br site until Ph 3.")
note(d, "GL COA comparison is fought on selection criteria drawn fm the deductions (surprise, mass at "
        "decisive pt, speed of build-up, protection of br site, simplicity). Course I wins four of five; "
        "Course II wins simplicity of dir keeping only. Decisive discriminator: Course II leaves VITARA "
        "(dominates BANANI/br site fm the north) uncleared until Ph 3 — doctrinally inadmissible once the br "
        "site is fixed at BANANI.")

heading(d, "SELECTION OF OWN BEST COURSE")
para(d, "34.\tAfter comparing merits and demerits against the en's most probable COA, I select Course I.")
note(d, "Stated as a comd's decision with the reasoning already argued — DSCSC style is decision-forward: "
        "the APRC has earned the one-line selection; the briefing (BM's brief Ser 6) carries the persuasion.")

heading(d, "OUTLINE PLAN")
para(d, "35.\tMsn.  8 Inf Bde will estb a brH across River DANUBE in gen area BAUNIA 7140 ASP to facilitate "
        "breakout of 9 Armd Bde.")
para(d, "36.\tExecution.  (Op Overlay Anx B; Gp Anx C.)")
para(d, "a.\tC of O.  Three phs as Course I; adv to waterway per para 28 con by CCO; bank gp (bn gp ex 7 "
        "Inf Bde) secures home bank and wins the dir-fire fight; one Cl 50 br site BANANI, raft sites S BANANI "
        "and AMBAGAN; ME: Ph 1 — 19 Inf Bn, Ph 2 — 39 Inf Bn, Ph 3 — 29 Inf Bn; defeat mechanism destruction "
        "(blind → suppress → strike → dislocate → destroy, per BM's brief Ser 7); res per para 22; risks: "
        "day mov to FUP (mitigated: ni mov, AD, CB) and frontal aslt (mitigated: FSP, tempo, darkness).",
     indent=0.3)
para(d, "b.\t19 Inf Bn.  Ph 1: aslt X, clr and secure VITARA; be prep asst 29 Bn at BAUNIA. Ph 2: be prep "
        "asst 39 Bn at SRIPUR. Ph 3: clr and secure ISLAMPUR (with JOLSHIRI shared as per bdry).", indent=0.3)
para(d, "c.\t29 Inf Bn.  Ph 1: aslt X, clr and secure BAUNIA; be prep asst 19 Bn at VITARA. Ph 2: be prep "
        "asst 39 Bn. Ph 3: clr and secure KHULNA.", indent=0.3)
para(d, "d.\t39 Inf Bn.  Ph 1: fol up, be prep asst fwd bns. Ph 2: aslt (ferried), clr and secure SRIPUR. "
        "Ph 3: be prep asst 19/29 Bn; find bde res coy.", indent=0.3)
para(d, "e.\tSqn 6 H.  X by hy raft fm A Hr; Ph 2 with 39 Bn; Ph 3 tps to 19/29 Bn; be prep defeat armd C "
        "attk fwd of PL TOP GUN.", indent=0.3)
para(d, "f.\tArty.  6 Med Regt: fire sp + CB. 7, 8 Fd and 9 SP Regt: DS aslt bns per FSP. MLRS Bty: CB + "
        "interdiction JOLSHIRI. SHORAD Bty: AD over X sites and br. SATA: acquire ISLAMPUR guns.", indent=0.3)
para(d, "g.\tEngr.  Two fd coy: entry/exit dev, breaching with aslt bns, then survivability; br coys (with "
        "CCO): ISR fm H+60, hy rafts fm A Hr, Cl 50 br fm Ph 2 secure. Pri: mob — survivability — counter-"
        "mob.", indent=0.3)
para(d, "h.\tRes.  One inf coy + sp coy (-) + fd pl; Comd: Coy Comd 39 Bn; loc BARISHAL (Ph 1) then BAUNIA; "
        "tasks: asst clearing, defeat C attk, on order.", indent=0.3)
para(d, "j.\tCoord Instr.  H Hr 132000 Jul; A Hr NB 140030; B Hr NB 140500; Bde Assy A DHAMRAI; CCO net "
        "active fm 131600; EW silence till H; deception: feint FERRY GHAT (live traffic + dummy boat pkts), "
        "dummy raft prep SAMPAN; smoke on LP/OPs H-15; PLs as overlay; Ack instr per SOP.", indent=0.3)
para(d, "37.\tSv Sp.  Fwd BAA: cen con, decen execution; ADS BOALI, sec BANANI; hel CASEVAC LZ BARISHAL; "
        "amn pri CB and smoke; boat/raft POL at BOLP; replen across riv on br open (emergency raft pkts by CCO "
        "pri).")
para(d, "38.\tComd and Sig.  Bde Tac with 29 Bn axis (far bank at A Hr); Main DHAMRAI; CCO net + line fwd "
        "of WCL; nicknames per X site; success codes by shuttle (red/green) until radio lifted at H.")
para(d, "Ack Instr:  Ack.", space_after=14)
para(d, "REZA", center=True, bold=True, space_after=0)
para(d, "Brig Gen", center=True, space_after=0)
para(d, "Comd 8 Inf Bde", center=True, space_after=8)
para(d, "Anx:", space_after=0)
para(d, "A.  Live Map Marking — RL Disposn (NATO/APP-6).", indent=0.3, space_after=0)
para(d, "B.  Op Overlay — BrH Op (Course I).", indent=0.3, space_after=0)
para(d, "C.  Sketch — X Area / CCO / Bank Master Gp Layout.", indent=0.3)

landscape_section(d); classification(d)
para(d, "ANX A TO APRC — LIVE MAP MARKING: RL DISPOSN", bold=True, underline=True, center=True, size=12)
picture(d, M1, width=9.6)
d.add_page_break()
para(d, "ANX B TO APRC — OP OVERLAY: BRH OP (COURSE I)", bold=True, underline=True, center=True, size=12)
picture(d, M2, width=9.6)
d.add_page_break()
para(d, "ANX C TO APRC — SKETCH: X AREA / CCO / BANK MASTER GP", bold=True, underline=True, center=True, size=12)
picture(d, M3, width=9.6)
d.save(os.path.join(S, "APRC_BrH_BAUNIA_Rewritten_with_Doctrinal_Notes.docx"))
print("doc2 saved")
